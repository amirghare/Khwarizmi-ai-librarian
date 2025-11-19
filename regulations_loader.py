from docx import Document
import os


class RegulationsLoader:
    def __init__(self, regulations_dir="data/regulations"):
        self.regulations_dir = regulations_dir
        self.regulations_text = ""
        self.load_all_regulations()

    def extract_text_from_docx(self, file_path):
        try:
            doc = Document(file_path)
            full_text = []

            # Extract all paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # Extract text from tables (if exists)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            full_text.append(text)

            return "\n".join(full_text)

        except Exception as e:
            print(f"❌ Error reading {file_path}: {e}")
            return ""

    def load_all_regulations(self):
        print("🔄 Loading library regulations...")

        if not os.path.exists(self.regulations_dir):
            print(f"⚠️ Directory {self.regulations_dir} does not exist!")
            return

        regulations_parts = []

        # List Word files
        docx_files = [f for f in os.listdir(self.regulations_dir) if f.endswith('.docx')]

        if not docx_files:
            print(f"⚠️ No Word files found in {self.regulations_dir}!")
            return

        # Load each file
        for filename in docx_files:
            file_path = os.path.join(self.regulations_dir, filename)
            print(f"   📄 Reading: {filename}")

            text = self.extract_text_from_docx(file_path)

            if text:
                # Add file title
                regulations_parts.append(f"\n{'='*60}")
                regulations_parts.append(f"📋 {filename}")
                regulations_parts.append(f"{'='*60}\n")
                regulations_parts.append(text)

        # Combine all texts
        self.regulations_text = "\n".join(regulations_parts)

        print(f"✅ {len(docx_files)} regulation files loaded")
        print(f"📊 Text volume: {len(self.regulations_text)} characters")

    def get_regulations_text(self):
        return self.regulations_text


if __name__ == "__main__":
    print("="*60)
    print("🧪 Testing RegulationsLoader")
    print("="*60)

    loader = RegulationsLoader()
    text = loader.get_regulations_text()

    if text:
        print(f"\n✅ Text loaded")
        print(f"📊 Length: {len(text)} characters")
        print(f"\n📄 Sample (first 200 characters):")
        print(text[:200])
    else:
        print("\n❌ No text loaded!")
